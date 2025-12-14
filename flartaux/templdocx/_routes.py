"""
Module docstring: This module provides functionality
for Cellmart wesite developed using flask.
"""
from flask import render_template, redirect, url_for, flash, request, current_app
from flask_sqlalchemy import SQLAlchemy
from db_file import db, login, login_manager, cache
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_caching import Cache
from templdocx import templdocx_bp
#, db, login_manager, cache

import sys
import platform
import flask
import re
import logging
import os
from app import db
from auth.forms import LoginForm, RegistrationForm
from models import TUser
from flask import render_template, redirect, url_for, flash, request 
from flask_login import current_user, login_user, logout_user 
import sqlalchemy as sa 
from urllib.parse import urlsplit
from flask import render_template
from flask_login import login_required
from forms import *
from forms import UserForm

from sqlalchemy import create_engine, MetaData
from sqlalchemy.orm import sessionmaker

#from db_init import db
from docxtpl import DocxTemplate
from flask import Response
from io import BytesIO
from flask import Flask, send_file

from datetime import datetime
from flask import Flask, request, redirect, url_for, session, flash
from flask import render_template
#from db_init import create_app
#from db_init import db
#from model import UserDetails, Product, OrderDetails, TCadastre, TParceldem, TCom2023, TCadastre, TPub, TUsager, TDemande, Country, City, Customer
from app import db
from models import TCadastre, TParceldem, TCom2023, TCadastre, TDemande
from models import TUser
from models import MainTable, SubTable
from templdocx import templdocx_bp

from sqlalchemy.sql import text
from sqlalchemy import desc, asc, and_, select, text
#from model import Base

import flask
import flask_login
from flask import Blueprint
from flask_login import LoginManager
from flask import Flask, render_template
from flask_htmx import HTMX
#Base = db.session.query

# Create the logger
logger = logging.getLogger('failed_login_attempts')
logger.setLevel(logging.INFO)

# Create a file handler
handler = logging.FileHandler('failed_login_attempts.log')
logger.addHandler(handler)

from flask import render_template
from flask.blueprints import Blueprint
from parsel import parsel_bp

from flask import Flask
import os.path

from sqlalchemy import Column, ForeignKey, Integer, String, create_engine, func, cast, Float 
from sqlalchemy.ext.declarative import declarative_base 
from sqlalchemy.orm import relationship,scoped_session,sessionmaker,aliased

from flask import Flask, render_template, request, jsonify, json
from flask_sqlalchemy import SQLAlchemy  
from wtforms import SelectField
from flask_wtf import FlaskForm

from flask import session
    
def setup_database(app):
    with app.app_context():
        db.create_all()

from werkzeug.debug import DebuggedApplication
'''
def create_app():
    # Insert whatever else you do in your Flask app factory.

    if app.debug:
        app.wsgi_app = DebuggedApplication(app.wsgi_app, evalex=True)

    return app
'''    
from flask_caching import Cache
import time

import sys
import platform
import flask
import re
import logging
import os
from datetime import datetime

#from database import db
from flask import Flask, request, redirect, url_for, session, flash
from flask import render_template
from flask.blueprints import Blueprint
from flask import current_app
from flask import Flask, session
from flask import request, flash
"""
dbinstance = Blueprint('dbinstance', __name__,
                 template_folder='templates',
                 static_folder='static')
"""                 
#from model import Base
import bcrypt

import os

import pandas as pd
from sqlalchemy import create_engine
from sqlalchemy.orm import relationship, Session
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.inspection import inspect
from sqlalchemy.dialects import postgresql
#from sqlalchemy import db.db.db.mapped_column, Table, Column, Bigdb.Integer, db.Integer, db.db.Text, db.Date, db.Boolean, db.db.db.String, ForeignKey
from sqlalchemy import desc
from flask_wtf import FlaskForm
from wtforms import StringField



# Function to generate the document
from datetime import datetime
from flask_wtf import FlaskForm
from wtforms import DateField
from wtforms.validators import DataRequired

from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_BREAK

from sqlalchemy import text

from sqlalchemy import Table, Column, Integer, Date
import datetime

from flask import render_template, redirect, url_for, session, send_file
import os
import tempfile
import time  
from flask import send_from_directory, request  


@templdocx_bp.route('/templdocx')
#@login_required
def index():
    return render_template('templdocx.html')


# Function to generate the document
def generate_docx(main_records, sub_records):
    start_time = time.time()  # Start timing the generation process    
    template_path = "templdocx/templates/template.docx"
    output_path = "static/generated_docs/output_with_page_breaks.docx"  # Save the generated file in a static folder
    
    # Create the document from the template
    doc = DocxTemplate(template_path)
    context = {"main_records": main_records, "sub_records": sub_records, "page_break": "\f"}
    doc.render(context)
    doc.save("static/generated_docs/output.docx")
    
    # Time after saving the first DOCX file
    elapsed_time = time.time() - start_time
    session['elapsed_time'] = elapsed_time  # Store elapsed time in session

    # Add page breaks (if necessary)
    document = Document("static/generated_docs/output.docx")
    for paragraph in document.paragraphs:
        if '\f' in paragraph.text:
            paragraph.text = paragraph.text.replace('\f', '')
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
    document.save(output_path)
    
    # Time after adding page breaks
    elapsed_time = time.time() - start_time
    session['elapsed_time'] = elapsed_time  # Update elapsed time in session
    
    # Save final DOCX with page breaks
    document.save("static/generated_docs/output_with_page_breaks.docx")
    elapsed_time = time.time() - start_time
    session['elapsed_time'] = elapsed_time  # Update elapsed time for final stage
    
    return elapsed_time  # Return the total time for generation

# Route to generate the document
@templdocx_bp.route('/templdocx/generate', methods=['POST'])
def generate():
    data = request.get_json()  # Receive the data from AJAX
    start_date = datetime.datetime.strptime(data['start_date'], '%Y-%m-%d').date()
    end_date = datetime.datetime.strptime(data['end_date'], '%Y-%m-%d').date()
    session['start_time'] = time.time()  # Store the start time in the session
    session['doc_generation_in_progress'] = True  # Set the generation flag to True

    # Sample start and end dates
    """
    start_date = datetime.date(2025, 1, 5)
    end_date = datetime.date(2025, 1, 6)
    """
    # Fetch main records from the database
    main_records_query = text("""
    SELECT DISTINCT no_interne, date_complet, demandeur, adrdem
    FROM 
    (SELECT DISTINCT no_interne, date_complet, u_nom_raison_sociale AS demandeur, (select distinct(concat(adr_postale_1, chr(10), adr_postalcp)) FROM t_usadresse
    RIGHT JOIN t_demande
    ON CAST(no_pacage_demandeur AS INT) = CAST(u_pacage AS INT)
    RIGHT JOIN t_usager
    ON CAST(no_pacage_demandeur AS INT) = CAST(adr_pacage AS INT)
    WHERE t_demande.no_interne = par_nointerne) AS adrdem FROM t_parceldem
    LEFT JOIN t_demande
    ON t_parceldem.par_nointerne = t_demande.no_interne             
    RIGHT JOIN t_usager 
    ON CAST(t_demande.no_pacage_demandeur AS INT) = CAST(t_usager.u_pacage AS INT)
    WHERE date_complet BETWEEN :start_date and :end_date
    AND no_interne LIKE :no_interne GROUP BY date_complet, par_nointerne,
    t_demande.no_interne, t_usager.u_pacage,
    u_nom_raison_sociale ORDER BY no_interne DESC) as bof 
    GROUP BY bof.no_interne, bof.date_complet, bof.demandeur, bof.adrdem
    ORDER BY bof.no_interne ASC
    """)
        

    params0 = {'no_interne': '%', 'start_date': start_date, 'end_date': end_date}
    result0 = db.session.execute(main_records_query, params0)
    main_rec = result0.fetchall()

    # Initialize main_records as a list to store records
    main_records = []
    for main_record in main_rec:
        main_records.append({
            'no_interne': main_record.no_interne,
            'date_complet': main_record.date_complet,
            'demandeur': main_record.demandeur,
            'adrdem': main_record.adrdem
        })

    # Prepare sub_records
    sub_records = {}
    for main_record in main_rec:
        # Now each `main_record` refers to an individual record
        params0 = {'no_interne': main_record.no_interne, 'start_date': start_date, 'end_date': end_date}

        sub_records_query = text("""
        SELECT DISTINCT libelle,         
        array_to_string(array_agg(CONCAT(w, z, x, y) ORDER BY z ASC, x ASC, y ASC), ' - ') 
        AS parcelles,
        SUM(CAST(par_surface AS INT)) AS superficie, 
        no_interne, date_complet 
        FROM 
        (SELECT DISTINCT libelle, CONCAT (SUBSTRING(t_parceldem.par_idsuf, 1, 2), SUBSTRING(t_parceldem.par_idsuf, 6, 3)), par_idsuf, par_surface,
        concat(t_parceldem.par_idsuf, ' ') AS u, SUBSTRING ( RIGHT( par_idsuf, 10 ) FROM '[A-Z]+') AS z,
        regexp_replace(SUBSTRING ( RIGHT( par_idsuf, 6 ) FROM '[0-9]+[_\\d]+' ), '(^|-)0*', '', 'g')::int AS x,
        SUBSTRING( par_idsuf, 6, 3 ) AS w, SUBSTRING ( RIGHT( par_idsuf, 6 ) FROM '[A-Z]+$' ) AS y,
        no_interne, date_complet 
        FROM t_parceldem
        LEFT JOIN t_demande
        ON t_parceldem.par_nointerne = t_demande.no_interne         
        RIGHT JOIN t_com2023
        ON com LIKE CONCAT(SUBSTRING(t_parceldem.par_idsuf, 1, 2), 
        CASE WHEN SUBSTRING(t_parceldem.par_idsuf, 6, 3) LIKE '000' THEN 
        SUBSTRING(t_parceldem.par_idsuf, 3, 3)
        ELSE CONCAT(SUBSTRING(t_parceldem.par_idsuf, 1, 2),SUBSTRING(t_parceldem.par_idsuf, 3, 3)) 
        END)
        AND libelle IN(SELECT libelle FROM t_com2023 
        WHERE dep = '22')
        OR com LIKE CONCAT(SUBSTRING(t_parceldem.par_idsuf, 1, 2),SUBSTRING(t_parceldem.par_idsuf, 3, 3)) 
        AND libelle IN(SELECT libelle FROM t_com2023 
        WHERE dep = '22')       
        RIGHT JOIN t_usager 
        ON CAST(t_demande.no_pacage_demandeur AS INT) = CAST(t_usager.u_pacage AS INT) 
        WHERE date_complet BETWEEN :start_date AND :end_date
        AND no_interne LIKE :no_interne 
        GROUP BY date_complet, par_nointerne, libelle,
        t_demande.no_interne, t_usager.u_pacage, t_parceldem.par_idsuf,
        t_parceldem.par_surface
        ORDER BY no_interne DESC) AS bof 
        GROUP BY bof.no_interne, bof.libelle, bof.date_complet
        ORDER BY bof.no_interne DESC, libelle ASC
        """)

        # Execute query and fetch all results
        result = db.session.execute(sub_records_query, {'no_interne': main_record.no_interne, 'start_date': start_date, 'end_date': end_date})

        # Fetch all rows from the query result
        sub_rec = result.fetchall()
        
        # Log the result to check the output before passing it to the template
        print("Fetched sub_records:", sub_rec)  # This is a debug print statement. You can use logging as well.


        # Store the sub-records for each main record in the sub_records dictionary
        sub_records[main_record.no_interne] = sub_rec

    
    # Debug the structure of sub_records:
    print("Sub records dictionary:", sub_records)
    
    # Generate the DOCX file
    elapsed_time = generate_docx(main_records, sub_records)

    session['doc_generation_in_progress'] = False  # Mark the process as complete
    return jsonify({
        'status': 'in_progress', 
        'elapsed_time': elapsed_time  # Return elapsed time to frontend for progress tracking
    })

    
# Route to check document generation status and get the elapsed time
@templdocx_bp.route('/templdocx/check_status')
def check_status():
    if session.get('doc_generation_in_progress', False):
        elapsed_time = time.time() - session['start_time']
        return jsonify({"status": "in_progress", "elapsed_time": elapsed_time})
    else:
        file_path = "static/generated_docs/output_with_page_breaks.docx"
        if os.path.exists(file_path):
            return jsonify({
                "status": "complete",
                "download_link": "/static/generated_docs/output_with_page_breaks.docx"
            })
        else:
            return jsonify({"status": "error", "message": "Document generation failed."})

# Route to serve the generated file for download
@templdocx_bp.route('/templdocx/send_file')
def send_file_route():
    file_path = "static/generated_docs/output_with_page_breaks.docx"
    if not os.path.exists(file_path):
        return "Error: No file found", 404
    return send_file(file_path, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')