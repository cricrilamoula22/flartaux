from flask import render_template
from flask import Flask, request, redirect, url_for, session, flash, jsonify, json
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user

from sqlalchemy.sql import text
from sqlalchemy import desc, asc, and_, select, text

from wtforms import SelectField
from flask_wtf import FlaskForm

from . import templdocx
#init_zut(app)
"""
import os
print(os.getcwd())
print(app.jinja_loader.searchpath)
"""
from ..database import db
from ..extensions import cache
#from .queries import fetch_dossiers_by_no_interne, fetch_dossiers_by_id, get_parcelle_by_idsuf, get_all_t_demande, get_all_parcelles, get_all_parceldem, get_all_t_com2023, fetch_sections_by_commune
#from .models import TParceldem, TCom2023, TCadastre
from .models import MainTable, SubTable
#from .forms import Form

import sys
import platform
import flask
import re
import logging
import os

import zipfile
"""
# Open the .docx file as a ZIP file
with zipfile.ZipFile('./static/generated_docs/output_with_page_breaks.docx', 'r') as docx:
    # List all the files inside the .docx
    print(docx.namelist())
"""

import sys
import platform
import flask
import re
import logging
import os
from flask import Blueprint, render_template, current_app, jsonify, json
from docx import Document
print(os.path.abspath('templates/templdocx/templates/template.docx'))
from flask import render_template, redirect, url_for, flash
from flask_login import login_user, logout_user, login_required
"""
#from . import templdocx_bp
from .models import MainTable, SubTable
#from .forms import Form
#from .queries import get_dossier
#from .methods import insertdoss
from ..database import db
"""
# Function to generate the document
from datetime import datetime
from flask_wtf import FlaskForm
from wtforms import DateField
from wtforms.validators import DataRequired

from docxtpl import DocxTemplate
from docx import Document
from docx.enum.text import WD_BREAK

from sqlalchemy import text

from sqlalchemy import Table, Column, Integer, Date
import datetime

from flask import render_template, redirect, url_for, session, send_file
import os
import tempfile
import time  
from flask import send_from_directory, request 

import zipfile

import threading
import time
from flask import send_from_directory

from datetime import timedelta

import datetime
import time
from flask import request, jsonify, session
import pypyodbc

@templdocx.route('/templdocx/testmanual')
def test_manual():
    from datetime import datetime
    from flask import jsonify

    # 📦 Données de test
    main_records = [
        {
            "no_interne": "2025-TEST001",
            "nom": "Dupont",
            "prenom": "Jean"
        }
    ]

    sub_records = {
        "2025-TEST001": [
            {"libelle": "Réunion équipe", "durée": "2h"},
            {"libelle": "Formation interne", "durée": "1h30"},
        ]
    }

    try:
        elapsed_time = generate_docx_individual(
            main_records,
            sub_records,
            datetime(2025, 4, 3),
            datetime(2025, 4, 16)
        )

        return jsonify({
            "status": "OK",
            "elapsed_time": elapsed_time,
            "download_link": "/templdocx/download_zip"
        })

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": str(e)
        })


@templdocx.route('/templdocx/test_docx_gen')
def test_docx_gen():
    from datetime import datetime
    main_records = [{"no_interne": "2025-TEST001", "nom": "Nom Test"}]
    sub_records = {
        "2025-TEST001": [
            {"libelle": "Test activité", "durée": "1h"},
            {"libelle": "Autre tâche", "durée": "2h"}
        ]
    }
    try:
        elapsed_time = generate_docx_individual(main_records, sub_records, datetime(2025,4,3), datetime(2025,4,16))
        return jsonify({
            'status': 'OK',
            'elapsed_time': elapsed_time,
            'download_link': '/templdocx/download_zip'
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


def zip_generated_docs(output_dir):
    zip_filename = os.path.join(output_dir, 'documents_fusionnes.zip')
    with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(output_dir):
            for file in files:
                if file.endswith('.docx') and file != 'documents_fusionnes.zip':
                    file_path = os.path.join(root, file)
                    arcname = os.path.relpath(file_path, output_dir)
                    zipf.write(file_path, arcname)
    print(f"Archive ZIP créée : {zip_filename}")
    return zip_filename

@templdocx.route('/templdocx')
#@login_required
def index():
    return render_template('templdocx.html')


@templdocx.route('/testdocx')
def testdocx():
    template_path = os.path.join(current_app.root_path, 'templates', 'templdocx', 'templates', 'template.docx')
    if not os.path.exists(template_path):
        return "Template file not found at " + template_path

    try:
        doc = Document(template_path)
        # Continue processing the document
    except Exception as e:
        return str(e)

    return render_template('testdocx.html')

def generate_docx_individual(main_records, sub_records, start_date, end_date):
    import zipfile
    import time
    from docxtpl import DocxTemplate
    from docx import Document
    from docx.enum.text import WD_BREAK

    template_path = os.path.join(current_app.root_path, 'templates', 'templdocx', 'templates', 'template.docx')
    output_dir = os.path.join(current_app.root_path, 'static', 'generated_docs')
    os.makedirs(output_dir, exist_ok=True)

    start_time = time.time()
    generated_files = []

    for record in main_records:
        no_interne = record['no_interne']

        # Filtrage ciblé des sub_records
        sub_records_raw = sub_records.get(no_interne, [])
        sub_records_cleaned = [dict(row._mapping) if not isinstance(row, dict) else row for row in sub_records_raw]

        context = {
            'main_records': [record],  # Liste avec 1 seul dict
            'sub_records': {no_interne: sub_records_cleaned},
            'start_date': start_date.strftime('%d/%m/%Y'),
            'end_date': end_date.strftime('%d/%m/%Y'),
            'page_break': '\f'
        }

        try:
            doc = DocxTemplate(template_path)
            doc.render(context)
            output_path = os.path.join(output_dir, f"{no_interne}.docx")
            doc.save(output_path)

            # Ajout des sauts de page
            document = Document(output_path)
            for paragraph in document.paragraphs:
                if '\f' in paragraph.text:
                    paragraph.text = paragraph.text.replace('\f', '')
                    run = paragraph.add_run()
                    run.add_break(WD_BREAK.PAGE)
            document.save(output_path)

            generated_files.append(output_path)

        except Exception as e:
            print(f"Erreur pour le dossier {no_interne} : {e}")

    # Création du fichier ZIP
    zip_filename = os.path.join(output_dir, 'documents_fusionnes.zip')
    with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file_path in generated_files:
            zipf.write(file_path, os.path.basename(file_path))
    generated_files.append(zip_filename)

    session['generated_files'] = generated_files
    elapsed_time = time.time() - start_time
    session['elapsed_time'] = elapsed_time

    return elapsed_time


@templdocx.route('/templdocx/download_zip')
def download_zip():
    import threading
    import time
    from flask import send_from_directory

    zip_dir = os.path.join(current_app.root_path, 'static', 'generated_docs')
    zip_filename = 'documents_fusionnes.zip'
    zip_path = os.path.join(zip_dir, zip_filename)

    def delete_all(paths):
        time.sleep(30)
        for path in paths:
            try:
                if os.path.exists(path):
                    os.remove(path)
                    print(f"Supprimé : {path}")
            except Exception as e:
                print(f"Erreur lors de la suppression : {e}")

    files_to_delete = session.get('generated_files', [])
    threading.Thread(target=delete_all, args=(files_to_delete,)).start()

    return send_from_directory(zip_dir, zip_filename, as_attachment=True)


# Function to generate the document
def generate_docx(main_records, sub_records, start_date, end_date):
    start_time = time.time()  # Start timing the generation process    
    print(f"Starting document generation at {datetime.datetime.now()}")
    print(f"Start Date: {start_date}, End Date: {end_date}")
    
    template_path = os.path.join(current_app.root_path, 'templates', 'templdocx', 'templates', 'template.docx')
    output_dir = os.path.join(current_app.root_path, 'static', 'generated_docs')
    
    # Ensure the output directory exists
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Directory {output_dir} was created.")
    else:
        print(f"Directory {output_dir} already exists.")
    
    output_path = os.path.join(output_dir, 'output_with_page_breaks.docx')
    print(f"Saving document to: {output_path}")
    
    # Log the context to debug
    context = {
        "name": "John Doe",
        "main_records": main_records,
        "sub_records": sub_records,
        "start_date": start_date.strftime("%Y-%m-%d"),
        "end_date": end_date.strftime("%Y-%m-%d"),
        "page_break": "\f"
    }
    print(f"Rendering DOCX with the following context: {context}")
    
    doc = DocxTemplate(template_path)
    
    try:
        doc.render(context)
        print("Document rendered successfully.")
    except Exception as e:
        print(f"Error rendering document: {e}")
    
    try:
        doc.save(output_path)
        print(f"Document saved successfully at {output_path}")
    except Exception as e:
        print(f"Error saving document: {e}")
    
    # Add page breaks
    print("Adding page breaks to the document...")
    document = Document(output_path)
    for paragraph in document.paragraphs:
        if '\f' in paragraph.text:
            paragraph.text = paragraph.text.replace('\f', '')
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
    document.save(output_path)
    print("Added page breaks and saved final docx.")
    
    elapsed_time = time.time() - start_time
    session['elapsed_time'] = elapsed_time  # Update elapsed time in session
    print(f"Document generated in {elapsed_time} seconds.")
    
    return elapsed_time  # Return the total time for generation


# Route to generate the document

#from your_module import generate_docx_individual  # Ton générateur DOCX

def parse_hfsql_date(date_str):
    """Convertit une date HFSQL AAAAMMJJ en datetime.date, ou None si invalide."""
    if not date_str or len(date_str) != 8 or not date_str.isdigit():
        return None
    try:
        return datetime.date(int(date_str[0:4]), int(date_str[4:6]), int(date_str[6:8]))
    except ValueError:
        return None
"""
from flask import Blueprint, request, jsonify, session
from datetime import datetime, timedelta
import time
import pypyodbc
from docx import Document
from io import BytesIO
import zipfile
import os
"""
#templdocx = Blueprint('templdocx', __name__)

@templdocx.route('/templdocx/generate', methods=['POST'])
def generate():
    try:
        data = request.get_json()
        start_date = datetime.strptime(data['start_date'], '%Y-%m-%d').date()
        end_date = datetime.strptime(data['end_date'], '%Y-%m-%d').date()
        next_day = end_date + timedelta(days=1)
        no_interne_filter = data.get('no_interne', '%')

        session['start_time'] = time.time()
        session['doc_generation_in_progress'] = True

        conn = pypyodbc.connect("DSN=DSN_ARTAUX")
        cursor = conn.cursor()

        # 🔹 Requête principale
        main_query = """
        SELECT DISTINCT
            t_demande.no_interne,
            t_demande.date_complet,
            t_parcel_statut.date_pub_pref,
            t_usager.u_nom_raison_sociale AS demandeur,
            CONCAT(t_usadresse.adr_postale_1, CHR(10), t_usadresse.adr_postalcp) AS adrdem,
            cedant.u_nom_raison_sociale AS cedant
        FROM t_demande
        LEFT JOIN t_parceldem ON t_parceldem.par_nointerne = t_demande.no_interne
        LEFT JOIN t_parcel_statut ON t_parcel_statut.no_interne = t_demande.no_interne
        LEFT JOIN t_usager ON CAST(t_demande.no_pacage_demandeur AS INT) = CAST(t_usager.u_pacage AS INT)
        LEFT JOIN t_usadresse ON CAST(t_demande.no_pacage_demandeur AS INT) = CAST(t_usadresse.adr_pacage AS INT)
        LEFT JOIN t_usager AS cedant ON CAST(t_demande.no_pacage_cedant AS INT) = CAST(cedant.u_pacage AS INT)
        WHERE t_demande.date_complet >= ?
          AND t_demande.date_complet <= ?
          AND t_parcel_statut.date_pub_pref = ?
          AND t_demande.no_interne LIKE ?
          AND t_demande.motif_ctrl NOT IN ('11','3','5')
          AND t_demande.code_statut LIKE '3'
          AND t_demande.code_avis LIKE 'P'
          AND t_demande.date_editdecis IS NULL
          AND t_demande.date_signature_decision IS NULL
          AND t_demande.code_decision IS NULL
        ORDER BY t_demande.no_interne ASC
        """

        cursor.execute(main_query, [start_date, end_date, next_day, no_interne_filter, None])
        main_rec = cursor.fetchall()

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for row in main_rec:
                no_interne = row[0]
                doc = Document()
                doc.add_heading(f'Demande {no_interne}', level=1)
                doc.add_paragraph(f"Date complète : {row[1]}")
                doc.add_paragraph(f"Date publication préfectorale : {row[2]}")
                doc.add_paragraph(f"Demandeur : {row[3]}")
                doc.add_paragraph(f"Adresse demandeur :\n{row[4]}")
                doc.add_paragraph(f"Cédant : {row[5]}")

                # 🔹 Sous-requête ateliers
                cursor.execute("""
                SELECT t_type_atelier.lib_atelier, t_demande_batiment.nb_atelier, t_unite.libunite
                FROM t_demande_batiment
                JOIN t_type_atelier ON t_demande_batiment.id_atelier = t_type_atelier.idt_type_atelier
                JOIN t_unite ON t_demande_batiment.unite = t_unite.idt_unite
                WHERE t_demande_batiment.id_interne = ?
                """, [no_interne])
                ateliers = cursor.fetchall()
                if ateliers:
                    doc.add_heading("Ateliers", level=2)
                    for a in ateliers:
                        doc.add_paragraph(f"{a[0]} — {a[1]} {a[2]}")

                # 🔹 Sous-requête parcelles
                cursor.execute("""
                SELECT par_idsuf, par_surface
                FROM t_parceldem
                WHERE par_nointerne = ?
                """, [no_interne])
                parcelles = cursor.fetchall()
                if parcelles:
                    doc.add_heading("Parcelles", level=2)
                    for p in parcelles:
                        doc.add_paragraph(f"{p[0]} — {p[1]} m²")

                # 🔹 Ajout au ZIP
                doc_buffer = BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                zipf.writestr(f'demande_{no_interne}.docx', doc_buffer.read())

        zip_buffer.seek(0)
        session['doc_generation_in_progress'] = False
        elapsed_time = round(time.time() - session['start_time'], 2)

        zip_path = os.path.join('static', 'generated_demandes.zip')
        with open(zip_path, 'wb') as f:
            f.write(zip_buffer.read())

        return jsonify({
            'status': 'complete',
            'elapsed_time': elapsed_time,
            'message': "La génération du document est terminée.",
            'download_link': '<a href="/static/generated_demandes.zip" target="_blank">Téléchargez votre document ici</a>'
        })

    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)})


    
# Route to check document generation status and get the elapsed time
@templdocx.route('/templdocx/check_status')
def check_status():
    if session.get('doc_generation_in_progress', False):
        elapsed_time = time.time() - session['start_time']
        return jsonify({"status": "in_progress", "elapsed_time": elapsed_time})
    else:
        file_path = "static/generated_docs/output_with_page_breaks.docx"
        if os.path.exists(file_path):
            return jsonify({
                "status": "complete",
                "download_link": "/static/generated_docs/output_with_page_breaks.docx"
            })
        else:
            return jsonify({"status": "error", "message": "Document generation failed."})

# Route to serve the generated file for download
@templdocx.route('/templdocx/send_file')
def send_file_route():
    file_path = "static/generated_docs/output_with_page_breaks.docx"
    if not os.path.exists(file_path):
        return "Error: No file found", 404
    return send_file(file_path, as_attachment=True, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
